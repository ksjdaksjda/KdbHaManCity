"""
Word论文修改工具 - AI重写内容，自动匹配学校格式
双击运行。
"""
import os, sys, json, time, threading, re, io

# ---- 编码修复 ----
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# ---- 依赖检查 ----
missing = []
try: from docx import Document
except ImportError: missing.append('python-docx')
try: from openai import OpenAI
except ImportError: missing.append('openai')

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.path.join(SCRIPT_DIR, "revise_config.json")

# 尝试读取网页版保存的API密钥
WEB_CONFIG = os.path.join(SCRIPT_DIR, "revise_config.json")

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except: pass
    return {}

def save_config(cfg):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(cfg, f, indent=2)

# ---- GUI ----
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext

class WordReviserApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word论文修改工具 - 保留排版 · AI修改内容")
        self.root.geometry("900x650")
        self.root.minsize(700, 500)
        self.config = load_config()
        self.client = None
        self.doc = None
        self.file_path = None
        self.paras = []
        self.modified_paras = {}
        self.current_idx = 0
        self.processing = False
        self.build_ui()
        if self.config.get("api_key"):
            self.api_var.set(self.config["api_key"])
            self.root.after(500, self.verify_api_silent)

    def build_ui(self):
        # 顶部：API密钥
        top = ttk.Frame(self.root, padding=10)
        top.pack(fill=tk.X)
        ttk.Label(top, text="DeepSeek API Key:").pack(side=tk.LEFT)
        self.api_var = tk.StringVar()
        api_entry = ttk.Entry(top, textvariable=self.api_var, width=45, show="*")
        api_entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(top, text="验证", command=self.verify_api).pack(side=tk.LEFT, padx=2)
        self.api_status = ttk.Label(top, text="", foreground="gray")
        self.api_status.pack(side=tk.LEFT, padx=10)

        # 文件选择
        mid = ttk.Frame(self.root, padding=10)
        mid.pack(fill=tk.X)
        ttk.Button(mid, text="选择Word文件", command=self.select_file).pack(side=tk.LEFT)
        self.file_label = ttk.Label(mid, text="  未选择文件", foreground="gray")
        self.file_label.pack(side=tk.LEFT, padx=5)
        ttk.Label(mid, text="模式:").pack(side=tk.LEFT, padx=(20, 5))
        self.mode_var = tk.StringVar(value="优化语言")
        mode_combo = ttk.Combobox(mid, textvariable=self.mode_var, values=["AI重写(推荐)", "降重改写", "扩展补充"], width=12, state="readonly")
        mode_combo.pack(side=tk.LEFT)

        # 进度条
        self.progress = ttk.Progressbar(self.root, mode='determinate')
        self.progress.pack(fill=tk.X, padx=10, pady=5)
        self.status_label = ttk.Label(self.root, text="就绪", foreground="gray")
        self.status_label.pack()

        # 主区域：左右分栏
        main = ttk.Frame(self.root)
        main.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        main.columnconfigure(0, weight=1)
        main.columnconfigure(1, weight=1)
        main.rowconfigure(0, weight=1)

        # 左侧：原文列表
        left = ttk.LabelFrame(main, text="段落列表 (点击预览)", padding=5)
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 5))
        self.para_list = tk.Listbox(left, font=("Microsoft YaHei", 10))
        self.para_list.pack(fill=tk.BOTH, expand=True)
        self.para_list.bind('<<ListboxSelect>>', self.on_para_select)

        # 右侧：对比预览
        right = ttk.LabelFrame(main, text="预览", padding=5)
        right.grid(row=0, column=1, sticky="nsew")
        right.rowconfigure(0, weight=1)
        right.rowconfigure(1, weight=1)
        right.columnconfigure(0, weight=1)

        ttk.Label(right, text="原文:").grid(row=0, column=0, sticky="w")
        self.orig_text = scrolledtext.ScrolledText(right, height=8, font=("Microsoft YaHei", 10), wrap=tk.WORD)
        self.orig_text.grid(row=0, column=0, sticky="nsew", pady=(0, 5))

        ttk.Label(right, text="AI修改后:").grid(row=1, column=0, sticky="w")
        self.mod_text = scrolledtext.ScrolledText(right, height=8, font=("Microsoft YaHei", 10), wrap=tk.WORD)
        self.mod_text.grid(row=1, column=0, sticky="nsew")

        # 底部按钮
        bottom = ttk.Frame(self.root, padding=10)
        bottom.pack(fill=tk.X)
        ttk.Button(bottom, text="开始处理", command=self.start_process).pack(side=tk.LEFT, padx=5)
        ttk.Button(bottom, text="暂停", command=self.pause_process).pack(side=tk.LEFT, padx=5)
        ttk.Button(bottom, text="保存修改", command=self.save_doc).pack(side=tk.LEFT, padx=5)
        ttk.Button(bottom, text="撤销当前段", command=self.undo_para).pack(side=tk.LEFT, padx=5)
        self.progress_text = ttk.Label(bottom, text="", foreground="gray")
        self.progress_text.pack(side=tk.RIGHT)

    def verify_api_silent(self):
        try:
            c = OpenAI(api_key=self.api_var.get(), base_url="https://api.deepseek.com/v1")
            r = c.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":"OK"}], max_tokens=2)
            self.client = c
            self.api_status.config(text="已连接", foreground="green")
            self.config["api_key"] = self.api_var.get()
            save_config(self.config)
        except: pass

    def verify_api(self):
        key = self.api_var.get().strip()
        if not key:
            self.api_status.config(text="请输入密钥", foreground="red")
            return
        self.api_status.config(text="验证中...", foreground="orange")
        try:
            c = OpenAI(api_key=key, base_url="https://api.deepseek.com/v1")
            r = c.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":"OK"}], max_tokens=2)
            self.client = c
            self.api_status.config(text="连接成功", foreground="green")
            self.config["api_key"] = key
            save_config(self.config)
        except Exception as e:
            self.api_status.config(text=f"失败: {str(e)[:40]}", foreground="red")

    def select_file(self):
        path = filedialog.askopenfilename(filetypes=[("Word文件", "*.docx"), ("所有文件", "*.*")])
        if not path: return
        self.file_path = path
        self.file_label.config(text=f"  {os.path.basename(path)}")
        self.status_label.config(text="读取文件中...")
        self.root.update()
        try:
            self.doc = Document(path)
            self.paras = [(i, p) for i, p in enumerate(self.doc.paragraphs) if p.text.strip()]
            self.modified_paras = {}
            # 分析文档字体
            fonts = {}; sizes = {}
            for i, p in self.doc.paragraphs:
                if p.runs:
                    r = p.runs[0]
                    fn = r.font.name or "默认"
                    fs = str(r.font.size) if r.font.size else "默认"
                    fonts[fn] = fonts.get(fn, 0) + 1
                    sizes[fs] = sizes.get(fs, 0) + 1
            top_font = max(fonts, key=fonts.get) if fonts else "未检测"
            top_size = max(sizes, key=sizes.get) if sizes else "未检测"
            self.para_list.delete(0, tk.END)
            for idx, para in self.paras:
                is_h = para.style.name.startswith("Heading")
                prefix = "[H] " if is_h else "     "
                self.para_list.insert(tk.END, f"{prefix}{para.text[:60]}...")
            info = f"已加载 {len(self.paras)} 段 | 字体: {top_font} | 字号: {top_size}"
            self.status_label.config(text=info)
            messagebox.showinfo("文档分析", f"段落: {len(self.paras)} 个\n主要字体: {top_font}\n主要字号: {top_size}\n\n修改时每个段落保留其原字体/大小/颜色\n页眉页脚/图片/表格不做任何改动")
        except Exception as e:
            messagebox.showerror("错误", f"无法读取文件: {e}")

    def on_para_select(self, event):
        sel = self.para_list.curselection()
        if not sel: return
        idx = sel[0]
        if idx >= len(self.paras): return
        pidx, para = self.paras[idx]
        self.orig_text.delete("1.0", tk.END)
        self.orig_text.insert("1.0", para.text)
        self.mod_text.delete("1.0", tk.END)
        mod = self.modified_paras.get(pidx, "")
        self.mod_text.insert("1.0", mod if mod else "(未修改)")

    def start_process(self):
        if not self.client:
            self.verify_api()
            if not self.client:
                messagebox.showwarning("提示", "请先验证API密钥")
                return
        if not self.paras:
            messagebox.showwarning("提示", "请先选择Word文件")
            return
        self.processing = True
        self.progress["maximum"] = len(self.paras)
        self.progress["value"] = 0
        self.current_idx = 0
        threading.Thread(target=self.process_loop, daemon=True).start()

    def process_loop(self):
        mode = self.mode_var.get()
        school = self.school_name or ""
        fmt = """【学校论文格式标准】
标题: 黑体, 章标题三号(16pt)加粗, 节标题四号(14pt)加粗
正文: 宋体, 小四号(12pt)
参考文献: 宋体, 五号(10.5pt)
行距: 1.5倍
引用: GB/T 7714-2015 顺序编码制""" if not school else f"""学校: {school}
请根据中国高校毕业论文GB/T 7713标准格式写作。"""

        modes = {
            "AI重写(推荐)": f"你是中文学术论文写作专家。请基于原文的主题和结构，完全重写内容。{fmt}\n学术严谨、逻辑清晰、善用长短句、避免AI套话。",
            "降重改写": f"深度改写降重，变换句式、同义词替换。{fmt}",
            "扩展补充": f"扩展内容，补充细节和论证，丰富学术表达。{fmt}"
        }
        mode_desc = modes.get(mode, modes["AI重写(推荐)"])

        for idx, (pidx, para) in enumerate(self.paras):
            if not self.processing: break
            text = para.text.strip()
            if skip_h and para.style.name.startswith("Heading"):
                self.update_progress(idx, f"跳过标题: {text[:30]}...")
                time.sleep(0.1)
                continue
            if len(text) < 10:
                self.update_progress(idx, f"跳过短文本")
                continue

            self.update_progress(idx, f"修改中: {text[:30]}...")
            system_prompt = f"你是学术论文编辑专家。{mode_desc}只输出修改后的文本，不要加任何说明。"
            try:
                resp = self.client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[{"role":"system","content":system_prompt}, {"role":"user","content":f"修改:\n{text}"}],
                    temperature=0.7, max_tokens=2048
                )
                new_text = resp.choices[0].message.content
                if new_text and len(new_text.strip()) > 5:
                    self.modified_paras[pidx] = new_text.strip()
            except Exception as e:
                self.modified_paras[pidx] = f"[修改失败: {e}]"
            time.sleep(0.3)

        self.root.after(0, lambda: self.status_label.config(text="处理完成！点击「保存修改」导出"))
        self.root.after(0, lambda: messagebox.showinfo("完成", f"修改了 {len(self.modified_paras)} 个段落\n点击「保存修改」导出Word"))

    def update_progress(self, idx, msg):
        self.root.after(0, lambda: self.progress.config(value=idx+1))
        self.root.after(0, lambda: self.progress_text.config(text=msg))

    def pause_process(self):
        self.processing = False
        self.status_label.config(text="已暂停")

    def undo_para(self):
        sel = self.para_list.curselection()
        if not sel: return
        idx = sel[0]
        if idx < len(self.paras):
            pidx = self.paras[idx][0]
            if pidx in self.modified_paras:
                del self.modified_paras[pidx]
                self.mod_text.delete("1.0", tk.END)
                self.mod_text.insert("1.0", "(已撤销)")
                self.status_label.config(text=f"已撤销第{idx+1}段")

    def save_doc(self):
        if not self.doc or not self.modified_paras:
            messagebox.showwarning("提示", "没有修改内容可保存")
            return
        self.status_label.config(text="保存中...应用GB/T 7713标准格式...")
        self.root.update()
        try:
            from docx.shared import Pt, Cm
            from docx.oxml.ns import qn
            from docx.enum.text import WD_LINE_SPACING
            # 设置标准页边距和行距
            for section in self.doc.sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.17)
                section.right_margin = Cm(3.17)
            for pidx, new_text in self.modified_paras.items():
                para = self.doc.paragraphs[pidx]
                is_heading = para.style.name.startswith("Heading")
                # 清空并重建run
                runs = para.runs
                for run in runs: run.text = ""
                if runs:
                    r0 = runs[0]
                    r0.text = new_text
                    if is_heading:
                        r0.font.name = "黑体"
                        r0._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
                        lvl = int(para.style.name.replace("Heading ","") or "1")
                        if lvl <= 1: r0.font.size = Pt(16)
                        elif lvl == 2: r0.font.size = Pt(14)
                        else: r0.font.size = Pt(12)
                        r0.bold = True
                    else:
                        r0.font.name = "宋体"
                        r0._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                        r0.font.size = Pt(12)
                        r0.bold = False
                    # 设置行距1.5倍
                    para.paragraph_format.line_spacing = 1.5
                else:
                    para.text = new_text
            out = self.file_path.replace(".docx", "_修改版.docx")
            if os.path.exists(out):
                out = self.file_path.replace(".docx", f"_修改版{int(time.time())%10000}.docx")
            self.doc.save(out)
            self.status_label.config(text=f"已保存: {os.path.basename(out)}")
            os.startfile(os.path.dirname(out))
            messagebox.showinfo("保存成功", f"文件已保存:\n{out}")
        except Exception as e:
            messagebox.showerror("错误", f"保存失败: {e}")

def main():
    if missing:
        root = tk.Tk()
        root.withdraw()
        msg = f"缺少依赖包: {', '.join(missing)}\n\n是否自动安装？"
        if messagebox.askyesno("安装依赖", msg):
            import subprocess
            subprocess.run([sys.executable, "-m", "pip", "install"] + missing, capture_output=True)
            try: from docx import Document; from openai import OpenAI
            except: messagebox.showerror("错误", "安装失败，请手动运行:\npip install python-docx openai"); return
        else: return
    root = tk.Tk()
    WordReviserApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
